#!/usr/bin/env python3
"""Convert the proposal markdown into a clean, rebrandable Word document."""

import re
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

NAVY = RGBColor(0x1B, 0x2A, 0x41)
ACCENT = RGBColor(0x2E, 0x6F, 0x9E)
GREY = RGBColor(0x55, 0x5F, 0x6D)
BODY_FONT = "Calibri"


def shade(cell, hex_fill):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(el)


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    tr_pr.append(el)


INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)")


def add_runs(paragraph, text):
    """Render **bold**, *italic* and `code` inside a paragraph."""
    for piece in INLINE.split(text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            run = paragraph.add_run(piece[2:-2])
            run.bold = True
        elif piece.startswith("`") and piece.endswith("`"):
            run = paragraph.add_run(piece[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
        elif piece.startswith("*") and piece.endswith("*") and len(piece) > 2:
            run = paragraph.add_run(piece[1:-1])
            run.italic = True
        else:
            paragraph.add_run(piece)


def style_document(doc):
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, colour, before, after in (
        ("Heading 1", 19, NAVY, 22, 8),
        ("Heading 2", 14, NAVY, 18, 6),
        ("Heading 3", 11.5, ACCENT, 13, 4),
    ):
        st = doc.styles[name]
        st.font.name = BODY_FONT
        st.font.size = Pt(size)
        st.font.color.rgb = colour
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True


def add_footer(doc):
    footer = doc.sections[0].footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("AIployee   ·   hakeem@aiployee.co.za   ·   Confidential")
    run.font.size = Pt(8)
    run.font.color.rgb = GREY
    run.font.name = BODY_FONT


def parse_table(lines, i):
    """Consume a markdown pipe table starting at lines[i]; return (rows, next_i)."""
    rows = []
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw = lines[i].strip().strip("|")
        cells = [c.strip() for c in raw.split("|")]
        if not all(re.fullmatch(r":?-{2,}:?", c) for c in cells if c):
            rows.append(cells)
        i += 1
    return rows, i


def render_table(doc, rows):
    width = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=width)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        if r_idx == 0:
            set_repeat_header(table.rows[0])
        for c_idx in range(width):
            text = row[c_idx] if c_idx < len(row) else ""
            cell = cells[c_idx]
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.space_before = Pt(2)
            add_runs(para, text)
            if r_idx == 0:
                shade(cell, "1B2A41")
                for run in para.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.size = Pt(9.5)
                if not para.runs:
                    para.add_run("")
            else:
                if r_idx % 2 == 0:
                    shade(cell, "F2F5F8")
                for run in para.runs:
                    run.font.size = Pt(9.5)
    doc.add_paragraph()


def build(md_path, out_path):
    lines = open(md_path, encoding="utf-8").read().split("\n")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    style_document(doc)
    add_footer(doc)

    i = 0
    first_heading_done = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped in ("---", "***", "___"):
            i += 1
            continue

        if stripped.startswith("|"):
            rows, i = parse_table(lines, i)
            if rows:
                render_table(doc, rows)
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            i += 1
            continue

        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            i += 1
            continue

        if stripped.startswith("# "):
            title = stripped[2:].strip()
            if not first_heading_done:
                para = doc.add_paragraph()
                para.paragraph_format.space_after = Pt(2)
                run = para.add_run("AIPLOYEE")
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = ACCENT
                run.font.name = BODY_FONT
                doc.add_heading(title, level=1)
                first_heading_done = True
            else:
                doc.add_heading(title, level=1)
            i += 1
            continue

        if stripped.startswith("> "):
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(0.6)
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(8)
            add_runs(para, stripped[2:].strip())
            for run in para.runs:
                run.italic = True
                run.font.color.rgb = GREY
            i += 1
            continue

        if re.match(r"^[-*+] ", stripped):
            para = doc.add_paragraph(style="List Bullet")
            para.paragraph_format.space_after = Pt(3)
            add_runs(para, stripped[2:].strip())
            i += 1
            continue

        if re.match(r"^\d+\. ", stripped):
            para = doc.add_paragraph(style="List Number")
            para.paragraph_format.space_after = Pt(3)
            add_runs(para, re.sub(r"^\d+\.\s*", "", stripped))
            i += 1
            continue

        # Paragraph: join soft-wrapped lines.
        buf = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if (not nxt or nxt.startswith(("|", "#", "> ", "---"))
                    or re.match(r"^([-*+] |\d+\. )", nxt)):
                break
            buf.append(nxt)
            i += 1

        text = " ".join(buf)
        # Metadata block on the cover ("**Prepared for:** ...") stays compact.
        para = doc.add_paragraph()
        if text.startswith("*AIployee") and text.endswith("*"):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_runs(para, text)
            for run in para.runs:
                run.font.color.rgb = GREY
                run.font.size = Pt(9)
        else:
            add_runs(para, text)

    doc.save(out_path)
    print(f"wrote {out_path}")


if __name__ == "__main__":
    build(sys.argv[1], sys.argv[2])
